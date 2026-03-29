import streamlit as st
import requests
import io
import os
import cv2
import tempfile
import pytesseract
from fpdf import FPDF
from docx import Document 

# --- STEP 1: DYNAMIC URL LOGIC ---
if "BACKEND_URL" in st.secrets:
    API_BASE_URL = st.secrets["BACKEND_URL"]
else:
    API_BASE_URL = "http://127.0.0.1:8000"
    
# 1. Page Configuration (Must be the first Streamlit command)
st.set_page_config(page_title="Legal AI Suite", page_icon="⚖️", layout="wide")

# Custom CSS for a SaaS-like Login and Dashboard
st.markdown("""
    <style>
    .main { background-color: #f8f9fa; }
    .stButton>button { width: 100%; border-radius: 8px; font-weight: bold; }
    .login-box { padding: 2rem; border-radius: 10px; background-color: white; box-shadow: 0 4px 6px rgba(0,0,0,0.1); }
    </style>
    """, unsafe_allow_html=True)

# 2. Session State Initialization
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "role" not in st.session_state:
    st.session_state.role = None
if "indexed" not in st.session_state:
    st.session_state.indexed = False

# -----------------------------------------
# UTILITY FUNCTIONS
# -----------------------------------------
def export_chat_to_word(messages):
    doc = Document()
    doc.add_heading('AI Legal Strategy & Timeline Report', 0)
    
    if not messages:
        doc.add_paragraph("No strategy session recorded yet.")
        
    for msg in messages:
        role = "Legal Team (User)" if msg["role"] == "user" else "AI Senior Counsel"
        doc.add_heading(role, level=2)
        doc.add_paragraph(msg["content"])
        doc.add_paragraph("-" * 40) 
        
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def convert_image_to_text(image_path, min_confidence=0.6):
    img = cv2.imread(image_path)
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
    enhanced_img = clahe.apply(gray)
    
    custom_config = r'--oem 3 --psm 6'
    result = pytesseract.image_to_data(enhanced_img if enhanced_img is not None else img, config=custom_config, output_type=pytesseract.Output.DICT)
    
    text = ' '.join([result['text'][i] for i in range(len(result['text'])) if int(result['conf'][i]) >= min_confidence * 100])
    return text

def convert_text_to_pdf(text, pdf_file):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font('Arial', size=12)
    lines = [text[i:i+95] for i in range(0, len(text), 95)]
    for line in lines:
        pdf.cell(200, 8, txt=line.encode('latin-1', 'replace').decode('latin-1'), ln=True, align='L')
    pdf.output(pdf_file)

# -----------------------------------------
# VIEW 1: THE LOGIN PAGE
# -----------------------------------------
def login_page():
    col1, col2, col3 = st.columns([1, 1.5, 1])
    
    with col2:
        st.markdown("<br><br><br>", unsafe_allow_html=True) 
        st.markdown("<div class='login-box'>", unsafe_allow_html=True)
        st.title("⚖️ Legal AI Suite")
        st.markdown("Secure Portal Access")
        
        role = st.selectbox("Select Access Level", ["Lawyer (Defense/Plaintiff)", "Judge (Presiding)"])
        username = st.text_input("Username", placeholder="e.g., harvey.specter")
        password = st.text_input("Password", type="password")
        
        if st.button("Secure Login"):
            if username and password: 
                st.session_state.logged_in = True
                st.session_state.role = "Lawyer" if "Lawyer" in role else "Judge"
                st.rerun() 
            else:
                st.error("Please enter credentials.")
        st.markdown("</div>", unsafe_allow_html=True)

# -----------------------------------------
# SHARED COMPONENT: UPLOAD SIDEBAR
# -----------------------------------------
def document_sidebar():
    with st.sidebar:
        st.header(f"🏛️ {st.session_state.role} Portal")
        st.write(f"Logged in as: **Admin**")
        if st.button("Log Out", type="secondary"):
            st.session_state.logged_in = False
            st.session_state.indexed = False
            st.rerun()
            
        st.divider()
        st.subheader("Case Management")
        uploaded_files = st.file_uploader("Upload Case Files (PDF)", type="pdf", accept_multiple_files=True)
        
        if st.button("Index to Secure Vault"):
            if uploaded_files:
                with st.spinner("Encrypting and Indexing..."):
                    files = [("files", (f.name, f.getvalue(), "application/pdf")) for f in uploaded_files]
                    # CHANGED: Added API_BASE_URL
                    response = requests.post(f"{API_BASE_URL}/upload-docs", files=files)
                    
                    if response.status_code == 200:
                        st.session_state.indexed = True
                        st.success("Case data secured and ready.")
                    else:
                        st.error("Vault Error.")

        if st.session_state.indexed and st.session_state.role == "Lawyer":
            st.divider()
            st.markdown("### 📥 Export Strategy")
            st.caption("Download your timeline and AI research to Microsoft Word.")
            if "messages" in st.session_state and len(st.session_state.messages) > 0:
                word_file = export_chat_to_word(st.session_state.messages)
                st.download_button(
                    label="📄 Download to Word (.docx)",
                    data=word_file,
                    file_name="Legal_Strategy_Report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

# -----------------------------------------
# VIEW 2: LAWYER DASHBOARD (WITH TABS)
# -----------------------------------------
def lawyer_dashboard():
    document_sidebar()
    st.title("💼 Counsel Strategy Dashboard")
    
    tab1, tab2 = st.tabs(["💬 Strategy & Analysis", "📸 Scan Physical Document (OCR)"])
    
    with tab1:
        if st.session_state.indexed:
            st.info("🎯 **Objective:** Find loopholes, build arguments, and protect the client.")
            
            if st.button("📅 Generate Case Timeline", use_container_width=True):
                with st.spinner("Scanning case files for dates and events..."):
                    # CHANGED: Added API_BASE_URL
                    res = requests.post(f"{API_BASE_URL}/timeline")
                    if res.status_code == 200:
                        timeline_data = res.json()["timeline"]
                        if "messages" not in st.session_state:
                            st.session_state.messages = []
                        timeline_msg = f"**Here is the chronological timeline of this case:**\n\n{timeline_data}"
                        st.session_state.messages.append({"role": "assistant", "content": timeline_msg})
                        st.rerun()
                    else:
                        st.error("Failed to generate timeline.")
            
            st.divider()

            if "messages" not in st.session_state:
                st.session_state.messages = []

            for message in st.session_state.messages:
                with st.chat_message(message["role"]):
                    st.markdown(message["content"])

            use_kanoon = st.toggle("🌐 Search Live Web (Indian Kanoon)")

            if prompt := st.chat_input("Ask a follow-up or formulate your legal inquiry..."):
                with st.chat_message("user"):
                    st.markdown(prompt)
                
                history_string = "\n".join([f"{m['role']}: {m['content']}" for m in st.session_state.messages])
                st.session_state.messages.append({"role": "user", "content": prompt})
                
                with st.chat_message("assistant"):
                    with st.spinner("Searching Kanoon..." if use_kanoon else "Analyzing case law..."):
                        payload = {
                            "question": prompt,
                            "mode": "lawyer",
                            "chat_history": history_string,
                            "search_web": use_kanoon
                        }
                        # CHANGED: Added API_BASE_URL
                        res = requests.post(f"{API_BASE_URL}/ask", json=payload)
                        
                        if res.status_code == 200:
                            answer = res.json()["answer"]
                            st.markdown(answer)
                            st.session_state.messages.append({"role": "assistant", "content": answer})
                        else:
                            st.error("Error communicating with the AI Vault.")
        else:
            st.warning("Please upload case files in the sidebar to begin building your strategy.")

    with tab2:
        st.header("📸 Digitize Hard Copy")
        st.caption("Upload a photo of a physical document. Extract the text and convert it to a PDF for the AI Vault.")
        
        uploaded_image = st.file_uploader("Upload an Image", type=["jpg", "jpeg", "png"])

        if uploaded_image is not None:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as temp_file:
                temp_file.write(uploaded_image.read())
                image_path = temp_file.name

            col1, col2 = st.columns(2)
            with col1:
                st.image(image_path, caption="Original Image", use_container_width=True)

            with st.spinner("Extracting text via OCR..."):
                extracted_text = convert_image_to_text(image_path)
            
            with col2:
                st.subheader("Extracted Text")
                st.text_area("You can edit this before saving:", extracted_text, height=300)

                if st.button("📄 Convert & Download PDF", type="primary"):
                    pdf_file = "scanned_document.pdf"
                    convert_text_to_pdf(extracted_text, pdf_file)

                    with open(pdf_file, "rb") as file:
                        st.download_button(
                            label="Save PDF to PC",
                            data=file,
                            file_name="scanned_document.pdf",
                            mime="application/pdf",
                            use_container_width=True
                        )
            os.remove(image_path)

# -----------------------------------------
# VIEW 3: JUDGE DASHBOARD
# -----------------------------------------
def judge_dashboard():
    document_sidebar()
    st.title("⚖️ Judicial Review Dashboard")
    
    if st.session_state.indexed:
        st.info("🔍 **Objective:** Neutral evaluation of facts against the Commercial Courts Act.")
        query = st.text_area("Enter review parameters:", placeholder="e.g., Does the evidence submitted on Page 4 meet the threshold for a summary judgment?")
        
        if st.button("Evaluate Evidence"):
            if query:
                with st.spinner("Reviewing statutes..."):
                    payload = {"question": query, "mode": "judge", "chat_history": "", "search_web": False}
                    # CHANGED: Added API_BASE_URL
                    res = requests.post(f"{API_BASE_URL}/ask", json=payload)
                    if res.status_code == 200:
                        st.markdown("### 📜 Judicial Opinion")
                        st.write(res.json()["answer"])
                    else:
                        st.error("Error communicating with the AI Vault.")
    else:
        st.warning("Please upload the docket files in the sidebar to begin your review.")

# -----------------------------------------
# MAIN APP ROUTER
# -----------------------------------------
if not st.session_state.logged_in:
    login_page()
else:
    if st.session_state.role == "Lawyer":
        lawyer_dashboard()
    elif st.session_state.role == "Judge":
        judge_dashboard()